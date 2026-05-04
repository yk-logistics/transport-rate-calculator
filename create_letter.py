from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'TH SarabunPSK'
font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
style.element.rPr.rFonts.set(qn('w:cs'), 'TH SarabunPSK')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_thai_font(run, size=16, bold=False, color=None):
    run.font.name = 'TH SarabunPSK'
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:cs'), 'TH SarabunPSK')
    r.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    if bold:
        r.rPr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
    if color:
        run.font.color.rgb = color

def add_line(text, size=16, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_thai_font(run, size, bold, color)
    return p

def add_mixed_line(parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
    """parts = list of (text, size, bold, color)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for text, size, bold, color in parts:
        run = p.add_run(text)
        set_thai_font(run, size, bold, color)
    return p

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
RED = RGBColor(0xC0, 0x39, 0x2B)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# === Header ===
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_after = Pt(2)
run = p_logo.add_run('บริษัท วาย.เค. ลอจิสติค จำกัด')
set_thai_font(run, size=26, bold=True, color=DARK_BLUE)

p_eng = doc.add_paragraph()
p_eng.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eng.paragraph_format.space_after = Pt(2)
p_eng.paragraph_format.space_before = Pt(0)
run = p_eng.add_run('Y.K. Logistics Co., Ltd.')
set_thai_font(run, size=14, bold=False, color=ACCENT_BLUE)

# Horizontal line
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_line.paragraph_format.space_after = Pt(12)
pPr = p_line._element.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B3A5C"/>'
    '</w:pBdr>'
)
pPr.append(pBdr)

# === Subject ===
add_mixed_line([
    ('เรื่อง', 16, True, DARK_BLUE),
    ('\t', 16, False, None),
    ('ขอความอนุเคราะห์ปรับอัตราค่าขนส่งตามการปรับขึ้นราคาน้ำมันเชื้อเพลิง', 16, True, BLACK),
], space_after=4)

add_mixed_line([
    ('เรียน', 16, True, DARK_BLUE),
    ('\t', 16, False, None),
    ('ลูกค้าผู้มีอุปการคุณทุกท่าน', 16, False, BLACK),
], space_after=12)

# === Body paragraph 1 ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.first_line_indent = Cm(1.25)

text_parts = [
    ('เนื่องด้วยราคาน้ำมันเชื้อเพลิงมีการปรับตัวสูงขึ้นอย่างต่อเนื่องและรวดเร็วในช่วงสัปดาห์ที่ผ่านมา '
     'ส่งผลกระทบโดยตรงต่อต้นทุนการให้บริการขนส่งของบริษัทฯ อย่างมีนัยสำคัญ '
     'โดยมีรายละเอียดการปรับขึ้นดังนี้', 16, False, BLACK),
]
for text, size, bold, color in text_parts:
    run = p.add_run(text)
    set_thai_font(run, size, bold, color)

# === Fuel Price Table ===
fuel_data = [
    ('วันที่', 'ปรับขึ้น (บาท/ลิตร)'),
    ('18 / 03 / 2569', '+0.50'),
    ('21 / 03 / 2569', '+0.70'),
    ('24 / 03 / 2569', '+1.80'),
    ('26 / 03 / 2569', '+6.00'),
    ('รวมสะสม', '+9.00'),
]

table = doc.add_table(rows=len(fuel_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

col_widths = [Cm(6), Cm(6)]
for row_idx, (col1, col2) in enumerate(fuel_data):
    row = table.rows[row_idx]
    for col_idx, (text, width) in enumerate(zip([col1, col2], col_widths)):
        cell = row.cells[col_idx]
        cell.width = width
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        is_header = (row_idx == 0)
        is_total = (row_idx == len(fuel_data) - 1)

        if is_header:
            bg_color = '1B3A5C'
            text_color = WHITE
            bold = True
        elif is_total:
            bg_color = 'E8F0FE'
            text_color = RED
            bold = True
        else:
            bg_color = 'FFFFFF' if row_idx % 2 == 1 else 'F5F7FA'
            text_color = BLACK
            bold = False

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
        run.text = text
        set_thai_font(run, size=16, bold=bold, color=text_color)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

# Table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# === Body paragraph 2 - rate info ===
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p2.paragraph_format.space_after = Pt(6)
p2.paragraph_format.first_line_indent = Cm(1.25)

parts2 = [
    ('ทั้งนี้ ราคาน้ำมัน ณ วันที่ 26/03/2569 จะอยู่ที่ ', 16, False, BLACK),
    ('38.94 บาท/ลิตร', 16, True, RED),
    (' ซึ่งอยู่ในช่วงราคา ', 16, False, BLACK),
    ('38.00–38.99 บาท/ลิตร', 16, True, BLACK),
    (' เปลี่ยนแปลงจากช่วงเดิม 29.00–29.99 บาท/ลิตร คิดเป็นการปรับข้ามช่วงราคาทั้งสิ้น ', 16, False, BLACK),
    ('9 ช่วง', 16, True, RED),
]
for text, size, bold, color in parts2:
    run = p2.add_run(text)
    set_thai_font(run, size, bold, color)

# === Body paragraph 3 - request ===
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p3.paragraph_format.space_after = Pt(12)
p3.paragraph_format.first_line_indent = Cm(1.25)

parts3 = [
    ('บริษัท ', 16, False, BLACK),
    ('วาย.เค. ลอจิสติค จำกัด', 16, True, DARK_BLUE),
    (' จึงใคร่ขอความอนุเคราะห์ปรับเพิ่มอัตราค่าขนส่ง ', 16, False, BLACK),
    ('18% จากอัตราค่าขนส่งเดิม', 16, True, RED),
    (' โดยคำนวณตามหลักเกณฑ์การปรับค่าขนส่งตามราคาน้ำมัน (9 ช่วง × 2% ต่อช่วง)', 16, False, BLACK),
]
for text, size, bold, color in parts3:
    run = p3.add_run(text)
    set_thai_font(run, size, bold, color)

# === Section: Future criteria ===
p_section = doc.add_paragraph()
p_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_section.paragraph_format.space_after = Pt(8)
p_section.paragraph_format.space_before = Pt(6)

pPr2 = p_section._element.get_or_add_pPr()
pBdr2 = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="2E75B6"/>'
    '</w:pBdr>'
)
pPr2.append(pBdr2)
run = p_section.add_run('หลักเกณฑ์การปรับค่าขนส่งตามราคาน้ำมัน (สำหรับอนาคต)')
set_thai_font(run, size=18, bold=True, color=DARK_BLUE)

# === Criteria description ===
p_criteria = doc.add_paragraph()
p_criteria.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_criteria.paragraph_format.space_after = Pt(6)
p_criteria.paragraph_format.first_line_indent = Cm(1.25)
parts_c = [
    ('บริษัทฯ ขอถือโอกาสนี้เรียนยืนยันหลักเกณฑ์ที่ใช้ร่วมกัน เพื่อความชัดเจนและเป็นธรรมแก่ทุกฝ่าย ดังนี้', 16, False, BLACK),
]
for text, size, bold, color in parts_c:
    run = p_criteria.add_run(text)
    set_thai_font(run, size, bold, color)

# Bullet points
bullets = [
    'ราคาน้ำมันเปลี่ยนแปลงทุก ๆ 1.00 บาท/ลิตร → ปรับค่าขนส่ง ±2%',
    'อ้างอิงตามช่วงราคาน้ำมัน ดังตารางตัวอย่างด้านล่าง',
    'หากในอนาคตราคาน้ำมันปรับลดลง บริษัทฯ จะดำเนินการปรับลดค่าขนส่งให้ตามหลักเกณฑ์เดียวกัน',
]
for b in bullets:
    p_b = doc.add_paragraph(style='List Bullet')
    p_b.paragraph_format.space_after = Pt(2)
    p_b.paragraph_format.left_indent = Cm(2)
    for run in p_b.runs:
        run.clear()
    p_b.clear()
    run = p_b.add_run(b)
    set_thai_font(run, size=16, bold=False, color=BLACK)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# === Reference Table ===
ref_data = [
    ('ช่วงราคาน้ำมัน (บาท/ลิตร)', 'การปรับค่าขนส่ง'),
    ('29.00 – 29.99', 'อัตราฐาน'),
    ('30.00 – 30.99', '+2%'),
    ('31.00 – 31.99', '+4%'),
    ('32.00 – 32.99', '+6%'),
    ('33.00 – 33.99', '+8%'),
    ('34.00 – 34.99', '+10%'),
    ('35.00 – 35.99', '+12%'),
    ('36.00 – 36.99', '+14%'),
    ('37.00 – 37.99', '+16%'),
    ('38.00 – 38.99', '+18%'),
]

table2 = doc.add_table(rows=len(ref_data), cols=2)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = True

for row_idx, (col1, col2) in enumerate(ref_data):
    row = table2.rows[row_idx]
    is_header = (row_idx == 0)
    is_current = (row_idx == len(ref_data) - 1)

    for col_idx, text in enumerate([col1, col2]):
        cell = row.cells[col_idx]
        cell.width = Cm(7)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if is_header:
            bg = '1B3A5C'
            tc = WHITE
            bold = True
        elif is_current:
            bg = 'FFF3E0'
            tc = RED
            bold = True
        else:
            bg = 'FFFFFF' if row_idx % 2 == 1 else 'F5F7FA'
            tc = BLACK
            bold = False

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

        run = cell.paragraphs[0].add_run(text)
        set_thai_font(run, size=15, bold=bold, color=tc)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

tbl2 = table2._tbl
tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders2 = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
    '</w:tblBorders>'
)
tblPr2.append(borders2)

# Note under table
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_note.paragraph_format.space_before = Pt(4)
p_note.paragraph_format.space_after = Pt(12)
run = p_note.add_run('* แถวสีส้มอ่อน คือช่วงราคาปัจจุบัน ณ วันที่ 26/03/2569')
set_thai_font(run, size=13, bold=False, color=RGBColor(0x99, 0x99, 0x99))
run.italic = True

# === Closing ===
p_close = doc.add_paragraph()
p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_close.paragraph_format.space_after = Pt(24)
p_close.paragraph_format.first_line_indent = Cm(1.25)
run = p_close.add_run(
    'บริษัทฯ หวังเป็นอย่างยิ่งว่าจะได้รับความเข้าใจและความร่วมมือจากท่าน '
    'และขอขอบพระคุณที่ไว้วางใจใช้บริการของบริษัทฯ ด้วยดีเสมอมา'
)
set_thai_font(run, size=16, bold=False, color=BLACK)

# === Signature ===
add_line('ขอแสดงความนับถือ', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=36)

p_sig_line = doc.add_paragraph()
p_sig_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sig_line.paragraph_format.space_after = Pt(4)
run = p_sig_line.add_run('___________________________')
set_thai_font(run, size=16, bold=False, color=BLACK)

add_line('บริษัท วาย.เค. ลอจิสติค จำกัด', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE, space_after=2)
add_line('Y.K. Logistics Co., Ltd.', size=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT_BLUE, space_after=0)

output_path = r'c:\Users\Home\Desktop\Project YK\หนังสือแจ้งปรับค่าขนส่ง.docx'
doc.save(output_path)
print('Saved successfully!')
